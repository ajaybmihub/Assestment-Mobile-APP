import os
import io
import fitz  # PyMuPDF
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import docx
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import logging
import cv2
import numpy as np

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def preprocess_image(pil_image):
    """
    Pre-processes image for better OCR results:
    1. Convert to Grayscale
    2. Thresholding (Binarization)
    3. Denoising
    """
    try:
        # Convert PIL image to OpenCV format
        open_cv_image = np.array(pil_image)
        # Convert RGB to BGR (OpenCV default) if necessary
        if len(open_cv_image.shape) == 3:
            open_cv_image = cv2.cvtColor(open_cv_image, cv2.COLOR_RGB2BGR)
        
        # 1. Grayscale
        gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
        
        # 2. Denoising
        denoised = cv2.fastNlMeansDenoising(gray, h=10)
        
        # 3. Thresholding (Otsu's Binarization)
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Convert back to PIL for Tesseract (Tesseract also accepts numpy arrays, but PIL is safer)
        return Image.fromarray(thresh)
    except Exception as e:
        logger.error(f"Image pre-processing failed: {e}")
        return pil_image # Fallback to original

app = FastAPI(title="Resume Extraction Service")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_from_pdf(file_bytes):
    """Extracts text from PDF. If text is sparse, uses OCR."""
    text = ""
    is_scanned = False
    
    try:
        # 1. Try direct extraction with PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = len(doc)
        for page in doc:
            text += page.get_text()
        
        # 2. Check if we need OCR (e.g., less than 100 chars across all pages)
        if len(text.strip()) < 100:
            logger.info("PDF appears to be scanned or empty. Starting OCR...")
            is_scanned = True
            text = ""
            # Convert PDF pages to images
            images = convert_from_bytes(file_bytes)
            page_count = len(images)
            for image in images:
                # Pre-process image for better OCR
                processed_image = preprocess_image(image)
                # Run OCR on processed image
                page_text = pytesseract.image_to_string(processed_image)
                text += page_text + "\n"
        
        doc.close()
        return text, is_scanned, page_count
    except Exception as e:
        logger.error(f"Error extracting from PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"PDF extraction failed: {str(e)}")

def extract_from_docx(file_bytes):
    """Extracts text from Word documents."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return "\n".join(full_text), False, 1
    except Exception as e:
        logger.error(f"Error extracting from DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"DOCX extraction failed: {str(e)}")

def extract_from_image(file_bytes):
    """Extracts text from Images using OCR."""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        # Pre-process image
        processed_image = preprocess_image(image)
        text = pytesseract.image_to_string(processed_image)
        return text, True, 1
    except Exception as e:
        logger.error(f"Error extracting from Image: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Image extraction failed: {str(e)}")

@app.post("/extract")
async def extract_resume(file: UploadFile = File(...)):
    """Main endpoint to handle multi-format resume extraction."""
    filename = file.filename
    content_type = file.content_type
    file_bytes = await file.read()
    
    logger.info(f"Received file: {filename}, Content-Type: {content_type}")
    
    try:
        if filename.lower().endswith(".pdf") or (content_type and content_type == "application/pdf"):
            text, is_scanned, pages = extract_from_pdf(file_bytes)
            format_type = "pdf"
        elif filename.lower().endswith((".docx", ".doc")) or (content_type and "officedocument" in content_type):
            text, is_scanned, pages = extract_from_docx(file_bytes)
            format_type = "docx"
        elif filename.lower().endswith((".png", ".jpg", ".jpeg")) or (content_type and content_type.startswith("image/")):
            text, is_scanned, pages = extract_from_image(file_bytes)
            format_type = "image"
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {content_type or 'unknown'}")
        
        return {
            "success": True,
            "text": text,
            "metadata": {
                "filename": filename,
                "format": format_type,
                "is_scanned": is_scanned,
                "page_count": pages
            }
        }
    except Exception as e:
        logger.error(f"Extraction error: {str(e)}")
        return {"success": False, "error": str(e)}

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
